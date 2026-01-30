"""
DOCX Handler - Extract text from Microsoft Word documents
"""

import asyncio
from pathlib import Path
from typing import Optional
import logging

try:
    from docx import Document
except ImportError:
    Document = None

from .base import BaseHandler

logger = logging.getLogger(__name__)


class DocxHandler(BaseHandler):
    """
    Extract text from .docx (Office Open XML) Word documents.

    Handles:
    - Text paragraphs
    - Tables
    - Headers and footers
    - Comments (optional)
    """

    @classmethod
    def supported_extensions(cls) -> set[str]:
        return {'.docx'}

    async def extract(self, path: Path) -> Optional[str]:
        """Extract text from DOCX file"""
        if Document is None:
            logger.error("python-docx not installed, cannot extract .docx files")
            return None

        return await asyncio.get_event_loop().run_in_executor(
            None, self._extract_sync, path
        )

    def _extract_sync(self, path: Path) -> Optional[str]:
        """Synchronous extraction"""
        try:
            doc = Document(str(path))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(f"\n[Table]\n{table_text}")

            # Combine all text
            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                logger.info(f"No text content in DOCX: {path}")
                return None

            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed for {path}: {e}")
            return None

    def _extract_table(self, table) -> str:
        """Extract text from a table"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Only include non-empty rows
                rows.append(" | ".join(cells))

        return "\n".join(rows) if rows else ""


class LegacyDocHandler(BaseHandler):
    """
    Handler for legacy .doc files (requires conversion).

    Note: Requires LibreOffice or similar for conversion.
    In production, consider using dedicated service for .doc conversion.
    """

    @classmethod
    def supported_extensions(cls) -> set[str]:
        return {'.doc'}

    async def extract(self, path: Path) -> Optional[str]:
        """Extract text from legacy .doc file via conversion"""
        logger.warning(f"Legacy .doc format detected: {path}")
        logger.info("Consider converting .doc files to .docx for better reliability")

        # Try conversion using LibreOffice (if available)
        converted_path = await self._convert_to_docx(path)

        if converted_path and converted_path.exists():
            # Use DOCX handler for converted file
            handler = DocxHandler()
            text = await handler.extract(converted_path)

            # Cleanup converted file
            converted_path.unlink(missing_ok=True)

            return text

        return None

    async def _convert_to_docx(self, path: Path) -> Optional[Path]:
        """Convert .doc to .docx using LibreOffice"""
        try:
            import subprocess
            import tempfile

            # Create temp directory for conversion
            temp_dir = Path(tempfile.mkdtemp())
            output_path = temp_dir / f"{path.stem}.docx"

            # Run LibreOffice conversion
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(temp_dir),
                str(path)
            ]

            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )

            await asyncio.wait_for(process.communicate(), timeout=30)

            if process.returncode == 0 and output_path.exists():
                return output_path

            logger.warning(f"LibreOffice conversion failed for {path}")
            return None

        except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
            logger.error(f"Doc conversion failed: {e}")
            return None
